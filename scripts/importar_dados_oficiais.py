"""Importa dados oficiais do repo CompStat-Rio para o sistema.

Fontes:
  - /tmp/compstat_dados/df_ocorrencias_tratado - Extração 1 .csv (115k linhas)
  - /tmp/compstat_dados/disk_denuncia.csv (83k linhas)
  - /tmp/compstat_dados/fatores_urbanos.csv (8k linhas)

Estratégia:
  1. Para cada CSV, percorre linha a linha
  2. Para cada registro com coordenada, identifica em qual polígono da FM caiu
     (via shapely.Polygon.contains)
  3. Mapeia campos para os schemas Pydantic
  4. Limita por área (top N mais recentes) para manter UI responsiva
  5. Substitui data/ocorrencias.json, data/denuncias.json, data/fatores_urbanos.json
"""

from __future__ import annotations

import csv
import json
import re
import sys
from datetime import datetime, date
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

DATA_DIR = ROOT / "data"
FONTE_DIR = Path("/tmp/compstat_dados")

CAP_OCORRENCIAS_POR_AREA = 200
CAP_DENUNCIAS_POR_AREA = 50


# Mapeamento desc_delito → TipoOcorrencia
MAP_DELITO = {
    "Roubo a transeunte": "roubo_transeunte",
    "Roubo de aparelho celular": "roubo_celular",
    "Roubo em coletivo": "roubo_coletivo",
    "Roubo de veículo": "roubo_veiculo",
    "Roubo a comércio": "roubo_comercio",
    "Furto a transeunte": "furto_transeunte",
    "Furto de aparelho celular": "furto_celular",
    "Furto de veículo": "furto_veiculo",
}

# Mapeamento tipo_ocorrencia_descricao → CategoriaFatorUrbano
MAP_CATEGORIA_FATOR = {
    "iluminação": "iluminacao_deficiente",
    "iluminacao": "iluminacao_deficiente",
    "vegetação": "vegetacao_obstrutiva",
    "vegetacao": "vegetacao_obstrutiva",
    "calçada": "calcada_obstruida",
    "calcada": "calcada_obstruida",
    "estacionamento": "estacionamento_irregular",
    "ônibus": "ponto_onibus_inseguro",
    "rua": "psr_concentrada",
    "população em situação": "psr_concentrada",
    "comércio": "comercio_irregular",
    "comercio": "comercio_irregular",
    "esconderijo": "esconderijo",
    "lixo": "lixo_entulho",
    "entulho": "lixo_entulho",
    "câmera": "ponto_cego_camera",
    "camera": "ponto_cego_camera",
}

# Mapeamento órgão (string oficial) → OrgaoMunicipal
MAP_ORGAO = {
    "COMLURB": "COMLURB",
    "SEOP": "SEOP",
    "CET-RIO": "CET_RIO",
    "CET_RIO": "CET_RIO",
    "RIOLUZ": "RIOLUZ",
    "SECONSERVA": "SECONSERVA",
    "SMAS": "SMAS",
    "SMS": "SMS",
    "GM-RIO": "GM_RIO",
    "GM_RIO": "GM_RIO",
}


def carregar_areas():
    """Carrega areas.json com geometrias para filtragem espacial."""
    from shapely import wkt
    areas_raw = json.loads((DATA_DIR / "areas.json").read_text(encoding="utf-8"))
    areas_geom = []
    for a in areas_raw:
        try:
            poly = wkt.loads(a["geometria_wkt"])
            # Pré-computa bbox para filtro rápido
            minx, miny, maxx, maxy = poly.bounds
            areas_geom.append({
                "poligono_id": a["poligono_id"],
                "nome": a["nome_area"],
                "poly": poly,
                "bbox": (minx, miny, maxx, maxy),
            })
        except Exception as e:
            print(f"⚠ falha em {a['poligono_id']}: {e}", file=sys.stderr)
    return areas_geom


def detectar_area(lat, lng, areas_geom):
    """Retorna poligono_id da área que contém o ponto, ou None."""
    from shapely.geometry import Point
    pt = Point(lng, lat)  # WKT: x=lng, y=lat
    for a in areas_geom:
        minx, miny, maxx, maxy = a["bbox"]
        if not (minx <= lng <= maxx and miny <= lat <= maxy):
            continue
        if a["poly"].contains(pt):
            return a["poligono_id"]
    return None


def parse_data_br(s: str):
    """Aceita 6/4/2020 8:16:00 ou 2024-05-20."""
    if not s:
        return None
    for fmt in ("%m/%d/%Y %H:%M:%S", "%m/%d/%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(s.strip(), fmt)
        except ValueError:
            continue
    return None


def inferir_dia_semana(dt: datetime):
    if not dt:
        return "seg"
    mapping = {0: "seg", 1: "ter", 2: "qua", 3: "qui", 4: "sex", 5: "sab", 6: "dom"}
    return mapping[dt.weekday()]


# ============================================================
# IMPORTAR OCORRÊNCIAS
# ============================================================

def importar_ocorrencias(areas_geom):
    arq = FONTE_DIR / "df_ocorrencias_tratado - Extração 1 .csv"
    if not arq.exists():
        print(f"⚠ não achei {arq}")
        return

    por_area: dict[str, list[dict]] = {}
    total_lidos = 0
    total_dentro = 0
    total_mapeados = 0

    with open(arq, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            total_lidos += 1
            try:
                lat = float(row["latitude"])
                lng = float(row["longitude"])
            except (ValueError, KeyError):
                continue

            pid = detectar_area(lat, lng, areas_geom)
            if not pid:
                continue
            total_dentro += 1

            tipo = MAP_DELITO.get(row.get("desc_delito", "").strip())
            if not tipo:
                continue
            total_mapeados += 1

            ano_str = row.get("ano", "")
            mes_str = row.get("mes", "1")
            hora_str = row.get("hora", "12")
            try:
                ano = int(ano_str) if ano_str else 2024
                mes = int(mes_str) if mes_str else 1
                hora = int(float(hora_str)) if hora_str else 12
                dia = 1  # CSV não traz dia
                dt = datetime(ano, mes, dia, hora, 0)
            except ValueError:
                dt = datetime(2024, 1, 1, 12, 0)

            dia_sem = row.get("dia_semana", "").strip().lower()
            if dia_sem in ("segunda", "segunda-feira"): dia_sem = "seg"
            elif dia_sem in ("terça", "terça-feira", "terca"): dia_sem = "ter"
            elif dia_sem in ("quarta", "quarta-feira"): dia_sem = "qua"
            elif dia_sem in ("quinta", "quinta-feira"): dia_sem = "qui"
            elif dia_sem in ("sexta", "sexta-feira"): dia_sem = "sex"
            elif dia_sem in ("sábado", "sabado"): dia_sem = "sab"
            elif dia_sem in ("domingo",): dia_sem = "dom"
            else: dia_sem = inferir_dia_semana(dt)

            por_area.setdefault(pid, []).append({
                "poligono_fm_id": pid,
                "tipo": tipo,
                "modalidade_crime": "a_pe",  # default, não vem no CSV
                "coordenada": {"lat": lat, "lng": lng, "descricao": None},
                "data_hora": dt.isoformat(),
                "dia_semana": dia_sem,
                "hora": hora,
                "descricao_modus": row.get("locf") or None,
            })

    # Limita por área
    saida: list[dict] = []
    for pid, items in por_area.items():
        items.sort(key=lambda x: x["data_hora"], reverse=True)
        saida.extend(items[:CAP_OCORRENCIAS_POR_AREA])

    (DATA_DIR / "ocorrencias.json").write_text(
        json.dumps(saida, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )
    print(f"[ocorrencias] {total_lidos} lidos | {total_dentro} dentro de polígonos "
          f"| {total_mapeados} com tipo mapeado | {len(saida)} salvos (cap {CAP_OCORRENCIAS_POR_AREA}/area)")
    for pid, items in por_area.items():
        print(f"  {pid}: {len(items)} → {min(len(items), CAP_OCORRENCIAS_POR_AREA)} mantidos")


# ============================================================
# IMPORTAR DENÚNCIAS
# ============================================================

def importar_denuncias(areas_geom):
    arq = FONTE_DIR / "disk_denuncia.csv"
    if not arq.exists():
        return

    por_area: dict[str, list[dict]] = {}
    total = 0

    with open(arq, encoding="latin-1") as f:
        reader = csv.DictReader(f, delimiter=";")
        for row in reader:
            total += 1
            lat_str = (row.get("latitude") or "").replace(",", ".")
            lng_str = (row.get("longitude") or "").replace(",", ".")
            try:
                lat = float(lat_str)
                lng = float(lng_str)
            except ValueError:
                continue

            pid = detectar_area(lat, lng, areas_geom)
            if not pid:
                continue

            dt = parse_data_br(row.get("data_denuncia", ""))
            if not dt:
                continue

            texto = (row.get("relato_redacted") or "").strip()[:1500]
            if not texto:
                continue

            tema = (row.get("classe") or row.get("tipos.tipo") or "").strip().lower()
            tema_curto = tema[:80] if tema else None

            por_area.setdefault(pid, []).append({
                "poligono_fm_id": pid,
                "data_recebimento": dt.isoformat(),
                "texto": texto,
                "local_mencionado": {"lat": lat, "lng": lng, "descricao": None},
                "horario_mencionado": None,
                "tema_principal": tema_curto,
            })

    saida: list[dict] = []
    for pid, items in por_area.items():
        items.sort(key=lambda x: x["data_recebimento"], reverse=True)
        saida.extend(items[:CAP_DENUNCIAS_POR_AREA])

    (DATA_DIR / "denuncias.json").write_text(
        json.dumps(saida, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )
    print(f"[denuncias] {total} lidas | {len(saida)} salvas")
    for pid, items in por_area.items():
        print(f"  {pid}: {len(items)} → {min(len(items), CAP_DENUNCIAS_POR_AREA)} mantidos")


# ============================================================
# IMPORTAR FATORES URBANOS
# ============================================================

def importar_fatores(areas_geom):
    arq = FONTE_DIR / "fatores_urbanos.csv"
    if not arq.exists():
        return

    por_area: dict[str, list[dict]] = {}
    total = 0

    with open(arq, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            total += 1
            # NOTA: no CSV oficial coordenada_x é LATITUDE e coordenada_y é LONGITUDE
            try:
                lat = float(row["coordenada_x"])
                lng = float(row["coordenada_y"])
            except (ValueError, KeyError):
                continue

            pid = detectar_area(lat, lng, areas_geom)
            if not pid:
                continue

            tipo_desc = (row.get("tipo_ocorrencia_descricao") or "").lower()
            categoria = "outro"
            for kw, cat in MAP_CATEGORIA_FATOR.items():
                if kw in tipo_desc:
                    categoria = cat
                    break

            orgao_raw = (row.get("orgao_responsavel") or "").upper().strip()
            orgao = MAP_ORGAO.get(orgao_raw, "FM")

            descricao = (row.get("tipo_ocorrencia_descricao") or "Fator urbano").strip()
            if row.get("logradouro"):
                descricao = f"{descricao} — {row['logradouro']}".strip()

            por_area.setdefault(pid, []).append({
                "poligono_fm_id": pid,
                "categoria": categoria,
                "descricao": descricao[:500],
                "coordenada": {"lat": lat, "lng": lng, "descricao": None},
                "orgao_responsavel": orgao,
                "severidade": "media",  # CSV não traz severidade explícita
            })

    saida: list[dict] = []
    for pid, items in por_area.items():
        saida.extend(items)

    (DATA_DIR / "fatores_urbanos.json").write_text(
        json.dumps(saida, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )
    print(f"[fatores] {total} lidos | {len(saida)} salvos")
    for pid, items in por_area.items():
        print(f"  {pid}: {len(items)} fatores")


# ============================================================
# MAIN
# ============================================================

# ============================================================
# IMPORTAR RELINTs (DOCX)
# ============================================================

# Mapeia padrão de nome de arquivo → poligono_id
MAP_RELINT_POLIGONO = {
    "Rodoviaria": "rodoviaria_terminal__02",
    "Metro_Botafogo": "metro_botafogo_rua_s_09",
    "Jardim_de_Alah": "jardim_de_alah_10",
    "Campo_Grande": "campo_grande_estacao_11",
    "Rio_Sul": "rio_sul_12",
    "Praia_Botafogo": "praia_de_botafogo_ru_14",
    "Estacoes_SFX": "estacoes_sao_francis_19",
    "Presidente_Vargas": "presidente_vargas_ca_20",
}

RE_HORARIO = re.compile(r"(\d{1,2})\s*[h:]?\s*(?:as|às|-|à)\s*(\d{1,2})\s*[h:]?", re.IGNORECASE)
RE_FACCAO = re.compile(r"\b(TCP|CV|ADA|milícia|milicia)\b", re.IGNORECASE)
RE_DIA = {
    "sex": re.compile(r"sexta", re.IGNORECASE),
    "sab": re.compile(r"sábado|sabado", re.IGNORECASE),
    "dom": re.compile(r"domingo", re.IGNORECASE),
    "seg": re.compile(r"segunda", re.IGNORECASE),
    "ter": re.compile(r"terça|terca", re.IGNORECASE),
    "qua": re.compile(r"quarta", re.IGNORECASE),
    "qui": re.compile(r"quinta", re.IGNORECASE),
}


def parse_relint_docx(path: Path, poligono_id: str, areas_geom) -> dict:
    """Extrai campos do RELINT a partir de um DOCX estruturado em tabela."""
    from docx import Document

    doc = Document(path)
    # Junta todas as células de todas as tabelas
    texto_completo = []
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texto_completo.append(cell.text.strip())
    # E parágrafos soltos
    for p in doc.paragraphs:
        if p.text.strip():
            texto_completo.append(p.text.strip())

    corpo = "\n".join(texto_completo)

    # Horário pico
    m = RE_HORARIO.search(corpo)
    horario_pico = (
        f"{int(m.group(1)):02d}:00-{int(m.group(2)):02d}:00"
        if m else "18:00-22:00"
    )

    # Dias críticos
    dias = [d for d, regex in RE_DIA.items() if regex.search(corpo)] or ["sex", "sab"]

    # Facção
    mfac = RE_FACCAO.search(corpo)
    faccao = None
    if mfac:
        s = mfac.group(1).upper()
        faccao = "milicia" if "MILIC" in s else s

    # Modus operandi: pega primeiros 1500 chars do corpo após o cabeçalho
    modus = corpo[:1500]

    # Centroide da área para coordenada de receptação/esconderijo placeholder
    centroide = next(
        (a for a in areas_geom if a["poligono_id"] == poligono_id), None,
    )
    if centroide:
        lat_c = centroide["poly"].centroid.y
        lng_c = centroide["poly"].centroid.x
    else:
        lat_c, lng_c = -22.9, -43.2

    # Citação: parágrafo curto
    citacao = corpo[:400].replace("\n", " ")

    return {
        "poligono_fm_id": poligono_id,
        "autor_orgao": "Secretaria-Geral CompStat",
        "data_documento": date.today().isoformat(),
        "modus_operandi_principal": modus,
        "modalidades_crime": ["a_pe"],
        "tipos_ocorrencia_alvo": ["roubo_transeunte", "roubo_celular"],
        "horario_pico": horario_pico,
        "dias_criticos": dias,
        "rotas_fuga": [],
        "pontos_receptacao": [],
        "esconderijos": [],
        "orcrim_influencia": faccao,
        "confianca": "alta",
        "citacao_fonte": citacao,
    }


def importar_relints(areas_geom):
    pasta = Path("/tmp/compstat_relints")
    if not pasta.exists():
        print("⚠ pasta de relints não encontrada")
        return

    relints = []
    for arq in sorted(pasta.glob("*.docx")):
        # Detecta poligono pelo nome do arquivo
        nome = arq.name
        poligono = None
        for chave, pid in MAP_RELINT_POLIGONO.items():
            if chave in nome:
                poligono = pid
                break
        if not poligono:
            print(f"  ⚠ {arq.name} sem mapeamento de área")
            continue

        try:
            rel = parse_relint_docx(arq, poligono, areas_geom)
            relints.append(rel)
            print(f"  ✓ {arq.name} → {poligono} (facção: {rel['orcrim_influencia']})")
        except Exception as e:
            print(f"  ✗ {arq.name}: {e}")

    # Backup
    import shutil
    origem = DATA_DIR / "relints.json"
    if origem.exists():
        backup = DATA_DIR / "relints.json.backup_sinteticos"
        if not backup.exists():
            shutil.copy2(origem, backup)
            print(f"  ⚙ backup: relints.json → {backup.name}")

    (DATA_DIR / "relints.json").write_text(
        json.dumps(relints, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )
    print(f"[relints] {len(relints)} salvos")


# ============================================================
# ENRIQUECER RELINTs COM DOMÍNIO TERRITORIAL
# ============================================================

def enriquecer_relints_com_dominio(areas_geom, raio_km: float = 3.0):
    """Cruza dominio_territorial.csv com áreas FM e atualiza facção dos RELINTs."""
    from shapely import wkt as shapely_wkt
    from shapely.geometry import Point

    arq = Path("/tmp/compstat_outros/dominio_territorial.csv")
    if not arq.exists():
        print("⚠ dominio_territorial.csv não baixado.")
        return

    # Coleta territórios de orcrim
    territorios = []
    with open(arq, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            try:
                poly = shapely_wkt.loads(row["geometria"])
                territorios.append({
                    "nome": row["nome_territorio"],
                    "orcrim": row["dominio_orcrim"].strip().upper(),
                    "poly": poly,
                    "centroide": (poly.centroid.y, poly.centroid.x),
                })
            except Exception:
                continue

    print(f"  {len(territorios)} territórios de orcrim carregados.")

    # Para cada área FM, busca facção dominante nas proximidades
    MAP_ORCRIM = {
        "TCP": "TCP", "CV": "CV", "ADA": "ADA",
        "MILICIA": "milicia", "MILÍCIA": "milicia",
        "FACÇÃO": "outras", "OUTRAS": "outras",
    }
    influencia_por_area: dict[str, str] = {}

    def dist_km(lat1, lng1, lat2, lng2):
        # Distância aproximada em km no Rio
        return ((lat1 - lat2) * 111) ** 2 + ((lng1 - lng2) * 102) ** 2

    for a in areas_geom:
        cy = a["poly"].centroid.y
        cx = a["poly"].centroid.x
        proximos = []
        for t in territorios:
            d2 = dist_km(cy, cx, *t["centroide"])
            if d2 < raio_km ** 2:
                proximos.append((d2, t))
        if proximos:
            proximos.sort()
            faccoes_perto = [
                MAP_ORCRIM.get(t["orcrim"], "outras")
                for _, t in proximos[:5]
            ]
            # Pega facção mais frequente
            from collections import Counter
            top_faccao = Counter(faccoes_perto).most_common(1)[0][0]
            influencia_por_area[a["poligono_id"]] = top_faccao
            print(f"  {a['poligono_id']:30s} ← {top_faccao} "
                  f"(de {len(proximos)} territórios próximos)")

    # Atualiza relints.json
    relints_path = DATA_DIR / "relints.json"
    if not relints_path.exists():
        return
    relints = json.loads(relints_path.read_text(encoding="utf-8"))
    for r in relints:
        if r["poligono_fm_id"] in influencia_por_area:
            r["orcrim_influencia"] = influencia_por_area[r["poligono_fm_id"]]

    relints_path.write_text(
        json.dumps(relints, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )
    print(f"  ✓ relints.json atualizado com facção territorial.")


# ============================================================
# IMPORTAR CPSR (POPULAÇÃO EM SITUAÇÃO DE RUA)
# ============================================================

def importar_cpsr_como_fatores(areas_geom):
    """Conta pessoas em situação de rua por área e adiciona como fator urbano."""
    try:
        import pandas as pd
    except ImportError:
        print("⚠ pandas não instalado, pulando CPSR")
        return

    arq = Path("/tmp/compstat_outros/CPSR.xlsx")
    if not arq.exists():
        return

    import warnings
    warnings.filterwarnings("ignore")
    df = pd.read_excel(arq)
    print(f"  CPSR: {len(df)} registros lidos")

    # Filtra com coordenadas
    if "Latitude" not in df.columns or "Longitude" not in df.columns:
        print("  ⚠ CPSR sem colunas Latitude/Longitude")
        return

    df = df.dropna(subset=["Latitude", "Longitude"])
    df = df[(df["Latitude"].between(-30, -20)) & (df["Longitude"].between(-50, -40))]
    print(f"  {len(df)} com coordenadas válidas")

    contagem_por_area: dict[str, int] = {}
    for _, row in df.iterrows():
        try:
            lat = float(row["Latitude"])
            lng = float(row["Longitude"])
        except (ValueError, TypeError):
            continue
        pid = detectar_area(lat, lng, areas_geom)
        if pid:
            contagem_por_area[pid] = contagem_por_area.get(pid, 0) + 1

    # Adiciona como fatores urbanos
    fatores_path = DATA_DIR / "fatores_urbanos.json"
    fatores = json.loads(fatores_path.read_text(encoding="utf-8"))

    n_novos = 0
    for pid, qtd in contagem_por_area.items():
        # Usa centroide da área como localização do fator agregado
        area = next(a for a in areas_geom if a["poligono_id"] == pid)
        cy = area["poly"].centroid.y
        cx = area["poly"].centroid.x

        severidade = (
            "critica" if qtd >= 30 else
            "alta" if qtd >= 10 else
            "media" if qtd >= 3 else "baixa"
        )
        fatores.append({
            "poligono_fm_id": pid,
            "categoria": "psr_concentrada",
            "descricao": (
                f"Censo CPSR (2020-2024): {qtd} pessoas em situação de rua "
                f"registradas na área."
            ),
            "coordenada": {"lat": cy, "lng": cx, "descricao": None},
            "orgao_responsavel": "SMAS",
            "severidade": severidade,
        })
        n_novos += 1
        print(f"  {pid}: {qtd} PSR ({severidade})")

    fatores_path.write_text(
        json.dumps(fatores, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )
    print(f"  ✓ {n_novos} fatores CPSR agregados.")


def main():
    print("[carga] áreas oficiais...")
    areas_geom = carregar_areas()
    print(f"  {len(areas_geom)} polígonos prontos para filtragem espacial.\n")

    # Backups
    import shutil
    for arq in ("ocorrencias.json", "denuncias.json", "fatores_urbanos.json"):
        origem = DATA_DIR / arq
        if origem.exists():
            backup = DATA_DIR / f"{arq}.backup_sinteticos"
            if not backup.exists():
                shutil.copy2(origem, backup)
                print(f"  ⚙ backup: {arq} → {backup.name}")

    print()
    importar_ocorrencias(areas_geom)
    print()
    importar_denuncias(areas_geom)
    print()
    importar_fatores(areas_geom)
    print()
    importar_relints(areas_geom)
    print()
    print("[enriquecimento] facção a partir de dominio_territorial")
    enriquecer_relints_com_dominio(areas_geom)
    print()
    print("[CPSR] população em situação de rua → fatores urbanos")
    importar_cpsr_como_fatores(areas_geom)


if __name__ == "__main__":
    main()
