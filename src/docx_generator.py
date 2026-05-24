"""
Gerador do Relatório Analítico de Área — formato oficial CompStat Municipal.

Estrutura conforme briefing técnico (Maio/2026), anexo p. 11-16:

  CAPA
  RESUMO EXECUTIVO (Score do bingo + tabela 4×4 de perguntas norteadoras)
  1. OCORRÊNCIAS CRIMINAIS
     - Identificação da Área
     - Indicadores do Período
     - Distribuição por Tipo (top 3)
     - Análise Temporal (heatmap 7×24 + período predominante + horário crítico)
  2. DINÂMICA CRIMINAL (síntese qualitativa)
  3. EFETIVO EMPREGADO – FORÇA MUNICIPAL (tabela 5×4)
  4. FATORES DE INCIDÊNCIA CRIMINAL (tabela por órgão + total câmeras)
  5. PLANO DE AÇÃO E RESPONSABILIZAÇÃO (pré-populado pela IA)

DOCX editável no Word/LibreOffice. Sem template externo.
"""

from __future__ import annotations

import io
from pathlib import Path
from datetime import datetime, date
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from schemas import (
    AreaPoligonoFM,
    BingoArea,
    Ocorrencia,
    RelintEstruturado,
    DenunciaDisque,
    FatorUrbano,
    RecomendacaoModalidade,
    QMD,
    ComparativoEvolucao,
    AcaoRecomendada,
)
from relatorio_compstat import (
    TOTAL_EFETIVO_FM,
    montar_identificacao,
    montar_indicadores,
    top_tipos_ocorrencia,
    montar_analise_temporal,
    montar_dinamica_criminal,
    montar_tabela_efetivo,
    gerar_perguntas_norteadoras,
    montar_tabela_fatores,
    montar_plano_acao,
)


# ============================================================
# CORES INSTITUCIONAIS
# ============================================================

COR_NAVY = RGBColor(0x1E, 0x27, 0x61)
COR_DESTAQUE = RGBColor(0xC6, 0x28, 0x28)
COR_OK = RGBColor(0x2E, 0x7D, 0x32)
COR_CINZA = RGBColor(0x55, 0x55, 0x55)


# ============================================================
# HELPERS DE TABELA
# ============================================================

def _bold_cell(cell, text: str, color: Optional[RGBColor] = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    if color:
        run.font.color.rgb = color


def _header_row(table, textos: list[str]) -> None:
    """Configura primeira linha da tabela como header (negrito + cinza claro)."""
    cells = table.rows[0].cells
    for c, t in zip(cells, textos):
        _bold_cell(c, t, COR_NAVY)


def _add_table(doc, n_rows: int, n_cols: int, style: str = "Light Grid"):
    table = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        table.style = style
    except KeyError:
        pass
    return table


# ============================================================
# GERAÇÃO PRINCIPAL
# ============================================================

def gerar_relatorio_docx(
    area: AreaPoligonoFM,
    bingo: BingoArea,
    recomendacao: RecomendacaoModalidade,
    qmd: QMD,
    *,
    ocorrencias_area: Optional[list[Ocorrencia]] = None,
    relints_area: Optional[list[RelintEstruturado]] = None,
    denuncias_area: Optional[list[DenunciaDisque]] = None,
    fatores_area: Optional[list[FatorUrbano]] = None,
    bingos_todos: Optional[list[BingoArea]] = None,
    comparativo_evolucao: Optional[ComparativoEvolucao] = None,
    acoes_outros_orgaos: Optional[list[AcaoRecomendada]] = None,
    heatmap_temporal_png: Optional[bytes] = None,
    grafico_evolucao_png: Optional[bytes] = None,
    mapa_segmentos_png: Optional[bytes] = None,
    efetivo_sugerido: Optional[int] = None,
    n_cameras: Optional[int] = None,
    periodo_inicio: Optional[date] = None,
    periodo_fim: Optional[date] = None,
    output_path: str = "output/relatorio.docx",
    # Compatibilidade com chamadas antigas
    padroes_disque=None,
) -> str:
    """Gera DOCX no formato oficial CompStat Municipal.

    Parâmetros obrigatórios: area, bingo, recomendacao, qmd.
    Demais são opcionais; se não fornecidos, seções omitem dados.
    """
    ocorrencias_area = ocorrencias_area or []
    relints_area = relints_area or []
    denuncias_area = denuncias_area or []
    fatores_area = fatores_area or []
    bingos_todos = bingos_todos or [bingo]
    efetivo_sugerido = efetivo_sugerido or qmd.efetivo_alocado

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============================================================
    # CAPA
    # ============================================================
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_inst.add_run("CompStat Municipal  |  Prefeitura do Rio de Janeiro")
    run.bold = True
    run.font.color.rgb = COR_NAVY
    run.font.size = Pt(10)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("RELATÓRIO ANALÍTICO DE ÁREA")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COR_NAVY

    subtit = doc.add_paragraph()
    subtit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtit.add_run("Subsídio para Reunião de CompStat")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = COR_CINZA

    doc.add_paragraph()

    # Bloco metadados da capa
    capa = _add_table(doc, 2, 2)
    capa.rows[0].cells[0].text = "Área de análise"
    capa.rows[0].cells[1].text = "Período de análise"
    for c in capa.rows[0].cells:
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = COR_NAVY

    capa.rows[1].cells[0].text = area.nome_area
    if periodo_inicio and periodo_fim:
        periodo_str = (
            f"Dados criminais e fatores ambientais: "
            f"{periodo_inicio:%d/%m/%Y} a {periodo_fim:%d/%m/%Y}"
        )
    else:
        periodo_str = "Todo o histórico disponível"
    capa.rows[1].cells[1].text = periodo_str

    doc.add_paragraph()

    # Total de roubos destacado
    n_roubos_capa = sum(
        1 for o in ocorrencias_area if o.tipo.startswith("roubo")
    )
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_total.add_run(
        f"Total de roubos no período: {n_roubos_capa} ocorrências"
    )
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COR_DESTAQUE

    # Mapa de segmentos quentes (se fornecido)
    if mapa_segmentos_png:
        doc.add_paragraph()
        p_map = doc.add_paragraph()
        p_map.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_map.add_run(f"Mapa de segmentos quentes — {area.nome_area}")
        run.bold = True
        run.font.size = Pt(10)
        doc.add_picture(io.BytesIO(mapa_segmentos_png), width=Inches(6.0))

    doc.add_page_break()

    # ============================================================
    # RESUMO EXECUTIVO
    # ============================================================
    doc.add_heading("RESUMO EXECUTIVO", level=1)

    # Caixa lateral com Score (nosso diferencial)
    box = _add_table(doc, 1, 3)
    box.rows[0].cells[0].text = "Score de risco"
    box.rows[0].cells[1].text = "Camadas ativas"
    box.rows[0].cells[2].text = "Bônus faccional"
    for c in box.rows[0].cells:
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = COR_NAVY

    doc.add_paragraph()
    cor = COR_DESTAQUE if bingo.componentes.score_final > 0.6 else COR_NAVY
    p_score = doc.add_paragraph()
    p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_score.add_run(
        f"{bingo.componentes.score_final:.2f}        "
        f"{bingo.n_camadas_ativas}/4        "
        f"x{bingo.componentes.bonus_faccional:.2f}"
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = cor

    doc.add_paragraph()

    # Tabela das 4 perguntas norteadoras
    analise_temporal = montar_analise_temporal(ocorrencias_area)
    perguntas = gerar_perguntas_norteadoras(
        area=area, bingo=bingo, recomendacao=recomendacao,
        ocorrencias_area=ocorrencias_area, relints_area=relints_area,
        fatores_area=fatores_area, analise_temporal=analise_temporal,
    )

    tab_pn = _add_table(doc, len(perguntas) + 1, 4)
    _header_row(tab_pn, [
        "Perguntas norteadoras",
        "Diagnóstico com base nos dados",
        "Operação FM / órgãos complementares",
        "Observações / sugestão (CompStat)",
    ])
    for i, p in enumerate(perguntas, 1):
        row = tab_pn.rows[i].cells
        row[0].text = p.pergunta
        row[1].text = p.diagnostico
        row[2].text = p.operacao
        row[3].text = p.observacao

    doc.add_page_break()

    # ============================================================
    # 1. OCORRÊNCIAS CRIMINAIS
    # ============================================================
    doc.add_heading("1. OCORRÊNCIAS CRIMINAIS", level=1)

    # 1.1 Identificação da Área
    doc.add_heading("Identificação da Área", level=2)
    ident = montar_identificacao(area, relints_area)
    items_ident = [
        ("Área FM", ident.nome_area),
        ("Número de trechos críticos", str(ident.n_trechos_criticos)),
        ("AISP", ident.aisp),
        ("Base FM", ident.base_fm),
        ("Bairro", ident.bairros),
        ("Subprefeitura", ident.subprefeitura),
        ("DP", ident.dp),
        ("BPM", ident.bpm),
        ("Área sob influência de grupo criminoso", ident.area_sob_influencia),
    ]
    tab_id = _add_table(doc, len(items_ident), 2)
    for i, (k, v) in enumerate(items_ident):
        _bold_cell(tab_id.rows[i].cells[0], k, COR_NAVY)
        tab_id.rows[i].cells[1].text = v

    doc.add_paragraph()

    # 1.2 Indicadores do Período
    doc.add_heading("Indicadores do Período", level=2)
    ind = montar_indicadores(
        ocorrencias_area, bingos_todos, area.poligono_id,
        comparativo_evolucao, periodo_inicio, periodo_fim,
    )
    tab_ind = _add_table(doc, 2, 6)
    _header_row(tab_ind, [
        "Período", "Roubos", "Furtos", "Total",
        "Ranking (%)", "Variação vs período anterior",
    ])
    row = tab_ind.rows[1].cells
    row[0].text = ind.periodo_label
    row[1].text = str(ind.roubos)
    row[2].text = str(ind.furtos)
    row[3].text = str(ind.total)
    row[4].text = f"{ind.posicao_ranking} ({ind.ranking_pct})"
    row[5].text = ind.variacao_anterior

    doc.add_paragraph()

    # 1.3 Distribuição por Tipo (top 3)
    doc.add_heading("Distribuição de ocorrências por tipo", level=2)
    top = top_tipos_ocorrencia(ocorrencias_area, n=3)
    if top:
        tab_top = _add_table(doc, len(top) + 1, 5)
        _header_row(tab_top, [
            "Ranking", "Tipo de ocorrência", "Qtd no período",
            "Data da última", "Variação",
        ])
        for i, t in enumerate(top, 1):
            row = tab_top.rows[i].cells
            row[0].text = t["rank"]
            row[1].text = t["tipo"]
            row[2].text = str(t["qtd"])
            row[3].text = t["data_ultima"]
            row[4].text = t["variacao"]
    else:
        doc.add_paragraph("Sem ocorrências no período analisado.")

    doc.add_paragraph()

    # 1.4 Análise Temporal
    doc.add_heading("Análise Temporal", level=2)
    if heatmap_temporal_png:
        doc.add_picture(io.BytesIO(heatmap_temporal_png), width=Inches(6.5))

    p_per = doc.add_paragraph()
    p_per.add_run("Período Predominante: ").bold = True
    p_per.add_run(analise_temporal.periodo_predominante)

    p_dia = doc.add_paragraph()
    p_dia.add_run("Dia / Horário Crítico: ").bold = True
    p_dia.add_run(analise_temporal.dia_horario_critico)

    doc.add_page_break()

    # ============================================================
    # 2. DINÂMICA CRIMINAL
    # ============================================================
    doc.add_heading("2. DINÂMICA CRIMINAL", level=1)
    doc.add_paragraph(
        "Fonte: Disque Denúncia e RELINTs da Força Municipal. "
        "Síntese qualitativa cruzada com a mancha quantitativa.",
        style="Intense Quote",
    )
    dinamica = montar_dinamica_criminal(relints_area, denuncias_area, ocorrencias_area)

    p1 = doc.add_paragraph()
    p1.add_run("Dinâmica do Crime: ").bold = True
    p1.add_run(dinamica["descricao"])

    p2 = doc.add_paragraph()
    p2.add_run("Modalidade Predominante: ").bold = True
    p2.add_run(dinamica["modalidade"])

    p3 = doc.add_paragraph()
    p3.add_run("Áreas de Fuga e Escoamento de Bens: ").bold = True
    p3.add_run(dinamica["areas_fuga"])

    if comparativo_evolucao:
        doc.add_paragraph()
        doc.add_heading("Evolução observada (período anterior vs. atual)", level=2)
        evo = _add_table(doc, 4, 3)
        _header_row(evo, [
            "Indicador",
            f"Antes ({comparativo_evolucao.snapshot_antes.data_referencia:%d/%m/%Y})",
            f"Depois ({comparativo_evolucao.snapshot_depois.data_referencia:%d/%m/%Y})",
        ])
        linhas_evo = [
            ("Roubos",
             comparativo_evolucao.snapshot_antes.total_roubos,
             comparativo_evolucao.snapshot_depois.total_roubos),
            ("Furtos",
             comparativo_evolucao.snapshot_antes.total_furtos,
             comparativo_evolucao.snapshot_depois.total_furtos),
            ("Score médio",
             f"{comparativo_evolucao.snapshot_antes.score_medio:.2f}",
             f"{comparativo_evolucao.snapshot_depois.score_medio:.2f}"),
        ]
        for i, (lab, a, b) in enumerate(linhas_evo, 1):
            evo.rows[i].cells[0].text = lab
            evo.rows[i].cells[1].text = str(a)
            evo.rows[i].cells[2].text = str(b)

        doc.add_paragraph(comparativo_evolucao.observacao)
        if grafico_evolucao_png:
            doc.add_picture(io.BytesIO(grafico_evolucao_png), width=Inches(6.5))

    doc.add_page_break()

    # ============================================================
    # 3. EFETIVO EMPREGADO – FORÇA MUNICIPAL
    # ============================================================
    doc.add_heading("3. EFETIVO EMPREGADO – FORÇA MUNICIPAL", level=1)
    doc.add_paragraph(
        f"Total municipal de referência: {TOTAL_EFETIVO_FM} agentes para "
        f"as 22 áreas prioritárias.",
        style="Intense Quote",
    )

    linhas_ef = montar_tabela_efetivo(area, recomendacao, bingo, efetivo_sugerido)
    tab_ef = _add_table(doc, len(linhas_ef) + 1, 4)
    _header_row(tab_ef, ["Campo", "Situação atual", "Sugestão", "Justificativa"])
    for i, le in enumerate(linhas_ef, 1):
        row = tab_ef.rows[i].cells
        row[0].text = le.campo
        row[1].text = le.situacao_atual
        row[2].text = le.sugestao
        row[3].text = le.justificativa

    doc.add_page_break()

    # ============================================================
    # 4. FATORES DE INCIDÊNCIA CRIMINAL
    # ============================================================
    doc.add_heading("4. FATORES DE INCIDÊNCIA CRIMINAL", level=1)
    doc.add_paragraph(
        "Fatores urbanos mapeados na área que contribuem para a "
        "incidência criminal, com órgão responsável pela intervenção.",
        style="Intense Quote",
    )

    linhas_fat = montar_tabela_fatores(fatores_area)
    if linhas_fat:
        tab_fat = _add_table(doc, len(linhas_fat) + 1, 3)
        _header_row(tab_fat, ["Fator identificado", "Descrição", "Responsável"])
        for i, lf in enumerate(linhas_fat, 1):
            row = tab_fat.rows[i].cells
            row[0].text = lf.fator
            row[1].text = lf.descricao
            row[2].text = lf.responsavel
    else:
        doc.add_paragraph("Sem fatores urbanos cadastrados para a área.")

    if n_cameras is not None:
        doc.add_paragraph()
        p_cam = doc.add_paragraph()
        p_cam.add_run("Câmeras identificadas na área: ").bold = True
        p_cam.add_run(f"{n_cameras} câmeras.")

    doc.add_page_break()

    # ============================================================
    # 5. PLANO DE AÇÃO E RESPONSABILIZAÇÃO
    # ============================================================
    doc.add_heading("5. PLANO DE AÇÃO E RESPONSABILIZAÇÃO", level=1)
    doc.add_paragraph(
        "Pré-populado pela plataforma a partir dos fatores urbanos e da "
        "recomendação operacional da FM. Deve ser revisto, ajustado e "
        "formalizado durante a reunião CompStat.",
        style="Intense Quote",
    )

    linhas_plano = montar_plano_acao(
        fatores_area, recomendacao, area, acoes_outros_orgaos,
    )
    if linhas_plano:
        tab_pl = _add_table(doc, len(linhas_plano) + 1, 4)
        _header_row(tab_pl, ["Ação acordada", "Responsável", "Prazo", "Status"])
        for i, la in enumerate(linhas_plano, 1):
            row = tab_pl.rows[i].cells
            row[0].text = la.acao
            row[1].text = la.responsavel
            row[2].text = la.prazo
            row[3].text = la.status
    else:
        doc.add_paragraph("Sem ações pré-populadas.")

    # Linhas adicionais em branco para preenchimento manual na reunião
    doc.add_paragraph()
    p_obs = doc.add_paragraph()
    run = p_obs.add_run(
        "Espaço para registro de compromissos adicionais assumidos na reunião:"
    )
    run.italic = True
    run.font.color.rgb = COR_CINZA
    tab_extra = _add_table(doc, 4, 4)
    _header_row(tab_extra, ["Ação acordada", "Responsável", "Prazo", "Status"])

    # ============================================================
    # RODAPÉ
    # ============================================================
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run(
        f"Secretaria-Geral do CompStat Municipal · "
        f"Gerado em {datetime.now():%d/%m/%Y %H:%M}"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = COR_CINZA

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ============================================================
# RELATÓRIO CONSOLIDADO (TODAS AS ÁREAS)
# ============================================================

def gerar_relatorio_geral_docx(
    por_area: dict,
    bingos_ranking: list[BingoArea],
    *,
    comparativos: Optional[list[ComparativoEvolucao]] = None,
    grafico_evolucao_png: Optional[bytes] = None,
    periodo_inicio: Optional[date] = None,
    periodo_fim: Optional[date] = None,
    output_path: str = "output/relatorio_consolidado.docx",
) -> str:
    """Gera DOCX consolidado cobrindo todas as áreas ativas da FM.

    Args:
        por_area: dict {poligono_id: {area, bingo, ocorrencias, relints,
                  denuncias, fatores, sugestao}}
        bingos_ranking: bingos ordenados (maior score primeiro)
        comparativos: lista de ComparativoEvolucao
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- CAPA CONSOLIDADA ----------
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_inst.add_run("CompStat Municipal  |  Prefeitura do Rio de Janeiro")
    run.bold = True
    run.font.color.rgb = COR_NAVY
    run.font.size = Pt(10)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("RELATÓRIO CONSOLIDADO")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = COR_NAVY

    subtit = doc.add_paragraph()
    subtit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtit.add_run(
        f"Todas as áreas prioritárias da Força Municipal ({len(por_area)} áreas)"
    )
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = COR_CINZA

    doc.add_paragraph()

    if periodo_inicio and periodo_fim:
        periodo_str = (
            f"Período: {periodo_inicio:%d/%m/%Y} a {periodo_fim:%d/%m/%Y}"
        )
    else:
        periodo_str = "Período: Todo o histórico disponível"

    p_per = doc.add_paragraph()
    p_per.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_per.add_run(periodo_str)
    run.bold = True
    run.font.size = Pt(11)

    # Totais
    n_oco_total = sum(len(d["ocorrencias"]) for d in por_area.values())
    n_relints_total = sum(len(d["relints"]) for d in por_area.values())
    n_fatores_total = sum(len(d["fatores"]) for d in por_area.values())

    doc.add_paragraph()
    totais = _add_table(doc, 2, 4)
    _header_row(totais, ["Áreas", "Ocorrências", "RELINTs", "Fatores"])
    totais.rows[1].cells[0].text = str(len(por_area))
    totais.rows[1].cells[1].text = str(n_oco_total)
    totais.rows[1].cells[2].text = str(n_relints_total)
    totais.rows[1].cells[3].text = str(n_fatores_total)

    doc.add_page_break()

    # ---------- RANKING ----------
    doc.add_heading("RANKING DAS ÁREAS POR SCORE DE RISCO", level=1)
    doc.add_paragraph(
        "Ordenação decrescente pelo score consolidado (mancha + RELINT + "
        "fator urbano + Disque + modus/rotas × bônus faccional).",
        style="Intense Quote",
    )

    rank_tab = _add_table(doc, len(bingos_ranking) + 1, 6)
    _header_row(rank_tab, [
        "Posição", "Área", "Score", "Camadas ativas",
        "Bônus faccional", "Facções",
    ])
    for i, b in enumerate(bingos_ranking, 1):
        row = rank_tab.rows[i].cells
        row[0].text = f"{i}º"
        row[1].text = b.nome_area
        row[2].text = f"{b.componentes.score_final:.2f}"
        row[3].text = f"{b.n_camadas_ativas}/4"
        row[4].text = f"x{b.componentes.bonus_faccional:.2f}"
        row[5].text = ", ".join(b.faccoes_envolvidas) if b.faccoes_envolvidas else "—"

    doc.add_page_break()

    # ---------- COMPARATIVO 90D ----------
    if comparativos:
        doc.add_heading("EVOLUÇÃO 90 DIAS — TODAS AS ÁREAS", level=1)
        doc.add_paragraph(
            "Comparativo de roubos e furtos antes vs depois da atuação da FM.",
            style="Intense Quote",
        )

        evo_tab = _add_table(doc, len(comparativos) + 1, 4)
        _header_row(evo_tab, [
            "Área", "Variação roubos", "Variação furtos", "Classificação",
        ])
        for i, c in enumerate(
            sorted(comparativos, key=lambda c: c.variacao_roubos_pct), 1,
        ):
            row = evo_tab.rows[i].cells
            row[0].text = c.nome_area
            row[1].text = f"{c.variacao_roubos_pct:+.1f}%"
            row[2].text = f"{c.variacao_furtos_pct:+.1f}%"
            row[3].text = c.classificacao.replace("_", " ")

        if grafico_evolucao_png:
            doc.add_paragraph()
            doc.add_picture(io.BytesIO(grafico_evolucao_png), width=Inches(6.5))

        doc.add_page_break()

    # ---------- FICHA RESUMIDA POR ÁREA ----------
    doc.add_heading("FICHAS RESUMIDAS POR ÁREA", level=1)

    # Ordena pelas mesmas posições do ranking
    ordem_areas = [b.poligono_fm_id for b in bingos_ranking]
    for pos, pid in enumerate(ordem_areas, 1):
        if pid not in por_area:
            continue
        d = por_area[pid]
        area = d["area"]
        bingo = d["bingo"]
        if bingo is None:
            continue

        doc.add_heading(f"{pos}º — {area.nome_area}", level=2)

        # Identificação compacta
        ident = montar_identificacao(area, d["relints"])
        ind = _ind_resumido(d["ocorrencias"], pid, bingos_ranking)

        info = _add_table(doc, 2, 4)
        info.rows[0].cells[0].text = "AISP"
        info.rows[0].cells[1].text = "Base FM"
        info.rows[0].cells[2].text = "Score"
        info.rows[0].cells[3].text = "Efetivo padrão"
        for c in info.rows[0].cells:
            for r in c.paragraphs[0].runs:
                r.bold = True
                r.font.color.rgb = COR_NAVY
        info.rows[1].cells[0].text = ident.aisp
        info.rows[1].cells[1].text = ident.base_fm
        info.rows[1].cells[2].text = f"{bingo.componentes.score_final:.2f}"
        info.rows[1].cells[3].text = f"{area.efetivo_padrao}"

        doc.add_paragraph()
        ind_tab = _add_table(doc, 2, 3)
        _header_row(ind_tab, ["Roubos", "Furtos", "Camadas ativas"])
        ind_tab.rows[1].cells[0].text = str(ind["roubos"])
        ind_tab.rows[1].cells[1].text = str(ind["furtos"])
        ind_tab.rows[1].cells[2].text = f"{bingo.n_camadas_ativas}/4"

        # Top 3 fatores urbanos
        if d["fatores"]:
            doc.add_paragraph()
            p_fat = doc.add_paragraph()
            p_fat.add_run("Top 3 fatores urbanos: ").bold = True
            fatores_grupo = montar_tabela_fatores(d["fatores"])[:3]
            p_fat.add_run("; ".join(
                f"{lf.fator} ({lf.responsavel})" for lf in fatores_grupo
            ))

        # Dinâmica resumida
        if d["relints"]:
            r0 = d["relints"][0]
            p_din = doc.add_paragraph()
            p_din.add_run("Dinâmica: ").bold = True
            p_din.add_run(r0.modus_operandi_principal[:300] + "...")

        doc.add_paragraph()

    doc.add_page_break()

    # ---------- PLANO DE AÇÃO CONSOLIDADO POR ÓRGÃO ----------
    doc.add_heading("PLANO DE AÇÃO CONSOLIDADO POR ÓRGÃO", level=1)
    doc.add_paragraph(
        "Ações agrupadas pelo órgão responsável, somando demandas de "
        "todas as áreas. Use como pauta única na reunião CompStat.",
        style="Intense Quote",
    )

    # Coleta todas as ações pré-populadas
    acoes_por_orgao: dict[str, list[tuple[str, str, str]]] = {}
    for pid, d in por_area.items():
        if not d["fatores"]:
            continue
        # Pega plano de ação da área (sem recomendacao da FM aqui)
        # Vamos só refletir os fatores
        fatores_grupo = montar_tabela_fatores(d["fatores"])
        for lf in fatores_grupo:
            acoes_por_orgao.setdefault(lf.responsavel, []).append((
                d["area"].nome_area, lf.fator, lf.descricao[:200],
            ))

    if acoes_por_orgao:
        for orgao, items in sorted(acoes_por_orgao.items()):
            doc.add_heading(f"{orgao} ({len(items)} demandas)", level=3)
            tab = _add_table(doc, len(items) + 1, 3)
            _header_row(tab, ["Área", "Fator", "Resumo"])
            for i, (area_nome, fator, desc) in enumerate(items, 1):
                row = tab.rows[i].cells
                row[0].text = area_nome
                row[1].text = fator
                row[2].text = desc
            doc.add_paragraph()

    # ---------- RODAPÉ ----------
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run(
        f"Secretaria-Geral do CompStat Municipal · "
        f"Gerado em {datetime.now():%d/%m/%Y %H:%M}"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = COR_CINZA

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _ind_resumido(
    ocorrencias_area: list[Ocorrencia],
    pid: str,
    bingos_todos: list[BingoArea],
) -> dict:
    """Indicadores resumidos para a ficha consolidada."""
    n_roubos = sum(1 for o in ocorrencias_area if o.tipo.startswith("roubo"))
    n_furtos = sum(1 for o in ocorrencias_area if o.tipo.startswith("furto"))
    return {"roubos": n_roubos, "furtos": n_furtos}
